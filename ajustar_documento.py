from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Abrir el documento
doc = Document('Nómina_Entrega_Útiles_2026_Completa_1-8.docx')

# Obtener la primera sección
section = doc.sections[0]

# Mostrar configuración actual
print(f"=== CONFIGURACIÓN ACTUAL ===")
print(f"Tamaño de página: {section.page_width.inches}\" x {section.page_height.inches}\"")
print(f"Márgenes - Top: {section.top_margin.inches}\", Bottom: {section.bottom_margin.inches}\", Left: {section.left_margin.inches}\", Right: {section.right_margin.inches}\"")
print(f"Orientación: {section.orientation}")
print(f"Número de párrafos: {len(doc.paragraphs)}")
print(f"Número de tablas: {len(doc.tables)}")

# Cambiar a orientación paisaje
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)  # Carta horizontal
section.page_height = Inches(8.5)

# Ajustar márgenes para paisaje (más pequeños para aprovechar el espacio)
section.top_margin = Inches(0.5)
section.bottom_margin = Inches(0.5)
section.left_margin = Inches(0.5)
section.right_margin = Inches(0.5)

print(f"\n=== NUEVA CONFIGURACIÓN ===")
print(f"Tamaño de página: {section.page_width.inches}\" x {section.page_height.inches}\"")
print(f"Márgenes - Top: {section.top_margin.inches}\", Bottom: {section.bottom_margin.inches}\", Left: {section.left_margin.inches}\", Right: {section.right_margin.inches}\"")
print(f"Orientación: {section.orientation}")

# Guardar el documento ajustado
doc.save('Nómina_Entrega_Útiles_2026_Completa_1-8.docx')
print("\n✅ Documento ajustado y guardado exitosamente")

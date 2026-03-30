from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# Abrir el documento
doc = Document('Nómina_Entrega_Útiles_2026_Completa_1-8.docx')

# Obtener la primera sección
section = doc.sections[0]

# Cambiar a orientación paisaje
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Inches(11)  # Carta horizontal
section.page_height = Inches(8.5)

# Márgenes muy pequeños para aprovechar espacio
section.top_margin = Inches(0.3)
section.bottom_margin = Inches(0.3)
section.left_margin = Inches(0.3)
section.right_margin = Inches(0.3)

# Ancho disponible después de márgenes (en twips)
page_width_twips = section.page_width
left_margin_twips = section.left_margin
right_margin_twips = section.right_margin
available_width_twips = page_width_twips - left_margin_twips - right_margin_twips
print(f"Ancho disponible: {available_width_twips / 1440}\" (en {available_width_twips} twips)")

# Ajustar todas las tablas
for table in doc.tables:
    # Establecer el ancho de la tabla al ancho disponible
    tblPr = table._element.tblPr
    if tblPr is None:
        tblPr = table._element.tblPr = doc._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblPr')
    
    # Crear elemento tblW si no existe
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = doc._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblW')
        tblPr.append(tblW)
    
    # Establecer ancho en twips (1/20 de punto)
    tblW.set(qn('w:w'), str(int(available_width_twips)))
    tblW.set(qn('w:type'), 'dxa')
    
    # Ajustar el ancho de cada celda proporcionalmente
    num_cols = len(table.rows[0].cells)
    cell_width = int(available_width_twips / num_cols)
    
    for row in table.rows:
        for cell in row.cells:
            tcPr = cell._element.tcPr
            if tcPr is None:
                tcPr = cell._element.tcPr = doc._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcPr')
            
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = doc._element.makeelement('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW')
                tcPr.append(tcW)
            
            tcW.set(qn('w:w'), str(cell_width))
            tcW.set(qn('w:type'), 'dxa')
    
    # Reducir tamaño de fuente en la tabla
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    if run.font.size is None or run.font.size > Pt(9):
                        run.font.size = Pt(9)

# Reducir también el tamaño de fuente en párrafos normales
for paragraph in doc.paragraphs:
    for run in paragraph.runs:
        if run.font.size and run.font.size > Pt(11):
            run.font.size = Pt(11)

print(f"\n=== CONFIGURACIÓN FINAL ===")
print(f"Tamaño de página: {section.page_width.inches}\" x {section.page_height.inches}\"")
print(f"Márgenes: {section.left_margin.inches}\" L, {section.right_margin.inches}\" R, {section.top_margin.inches}\" T, {section.bottom_margin.inches}\" B")
print(f"Orientación: LANDSCAPE")
print(f"Tablas ajustadas: {len(doc.tables)}")

# Guardar el documento
doc.save('Nómina_Entrega_Útiles_2026_Completa_1-8.docx')
print("\n✅ Documento completamente ajustado y guardado")
